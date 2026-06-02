from __future__ import annotations

import os
import tempfile
from pathlib import Path
from typing import List

from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
import uvicorn

try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

AI_IMPORT_ERROR = None
try:
    import requests
    from bs4 import BeautifulSoup
    import pypdf
    import docx
    import chromadb
    from google.cloud import speech
    from llama_index.core import VectorStoreIndex, StorageContext, Settings
    from llama_index.vector_stores.chroma import ChromaVectorStore
    from llama_index.llms.groq import Groq
    from llama_index.core.chat_engine import CondensePlusContextChatEngine
    from llama_index.embeddings.huggingface import HuggingFaceEmbedding
    from llama_index.core.node_parser import TokenTextSplitter
    from llama_index.core.schema import Document
except ImportError as exc:
    AI_IMPORT_ERROR = exc

# =====================================================================
# 1. INITIALIZATION & STORAGE ENGINE SETUP
# =====================================================================
DB_PATH = os.getenv("DATABASE_PATH", "./chroma_db") 
os.makedirs(DB_PATH, exist_ok=True)
PROJECT_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_WORKSPACE = os.getenv("DEFAULT_WORKSPACE", "second_brain")


GROQ_API_KEY = os.getenv("GROQ_API_KEY")
if AI_IMPORT_ERROR is None:
    if GROQ_API_KEY:
        Settings.llm = Groq(
            model="llama-3.3-70b-versatile",
            api_key=GROQ_API_KEY,
            temperature=0.2
        )

    Settings.embed_model = HuggingFaceEmbedding(model_name="all-MiniLM-L6-v2")
    chroma_client = chromadb.PersistentClient(path=DB_PATH)
else:
    chroma_client = None


GCP_KEY_PATH = os.getenv("GOOGLE_APPLICATION_CREDENTIALS", "gcp_key.json")
app = FastAPI(title="Second Brain AI Layer", version="1.0.0")
app.add_middleware(
    CORSMiddleware,
    allow_origins=os.getenv("CORS_ORIGINS", "*").split(","),
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class ChatRequest(BaseModel):
    message: str
    workspace_id: str = DEFAULT_WORKSPACE
    app_identity: str = "Second Brain"
    custom_rules: str = "Answer concisely and cite indexed memories when possible."


class UrlIngestRequest(BaseModel):
    url: str
    workspace_id: str = DEFAULT_WORKSPACE


class TextIngestRequest(BaseModel):
    text: str
    source_id: str = "manual-note"
    source_type: str = "text"
    workspace_id: str = DEFAULT_WORKSPACE


def require_ai_dependencies():
    if AI_IMPORT_ERROR is not None:
        raise RuntimeError(
            f"AI dependencies are not installed ({AI_IMPORT_ERROR}). "
            "Run: pip install -r requirements.txt"
        )

if not os.path.exists(GCP_KEY_PATH):
    print(f"⚠️ Warning: GCP credentials file missing at '{GCP_KEY_PATH}'.")

# =====================================================================
# 2. FILE INGESTION PARSERS (Multi-Modal Data Layer)
# =====================================================================
def extract_text_from_pdf(file_path: str) -> str:
    """Extracts raw text content from a local PDF file."""
    require_ai_dependencies()
    try:
        text = ""
        with open(file_path, "rb") as f:
            reader = pypdf.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text.strip()
    except Exception as e:
        raise RuntimeError(f"Failed to parse PDF file: {str(e)}")

def extract_text_from_docx(file_path: str) -> str:
    """Extracts raw text content from a local Word document (.docx)."""
    require_ai_dependencies()
    try:
        doc = docx.Document(file_path)
        full_text = [paragraph.text for paragraph in doc.paragraphs]
        return "\n".join(full_text).strip()
    except Exception as e:
        raise RuntimeError(f"Failed to parse DOCX file: {str(e)}")

def scrape_url_to_markdown(url: str) -> str:
    """Fetches a webpage, cleans non-content tags, and structures pure text."""
    require_ai_dependencies()
    try:
        headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"}
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.text, "html.parser")
        for element in soup(["script", "style", "nav", "footer", "header", "aside"]):
            element.decompose()
            
        lines = (line.strip() for line in soup.get_text().splitlines())
        chunks = (phrase for line in lines for phrase in line.split("  "))
        clean_text = "\n".join(chunk for chunk in chunks if chunk)
        
        return f"# Source: {url}\n\n{clean_text}"
    except Exception as e:
        raise RuntimeError(f"BeautifulSoup scraping failed: {str(e)}")

def transcribe_audio_gcp(audio_file_path: str) -> str:
    """Transcribes an input audio file via Google Cloud Speech-to-Text API."""
    require_ai_dependencies()
    try:
        client = speech.SpeechClient()
        with open(audio_file_path, "rb") as audio_file:
            content = audio_file.read()

        audio = speech.RecognitionAudio(content=content)
        config = speech.RecognitionConfig(
            encoding=speech.RecognitionConfig.AudioEncoding.MP3,
            sample_rate_hertz=16000,
            language_code="en-US",
        )

        response = client.recognize(config=config, audio=audio)
        transcript = " ".join([result.alternatives[0].transcript for result in response.results])
        return transcript if transcript else "[Empty Transcription]"
    except Exception as e:
        raise RuntimeError(f"GCP Speech-to-Text transcription failed: {str(e)}")

def extract_text_from_upload(file_path: str, filename: str) -> tuple[str, str]:
    """Routes uploaded files to the parser that matches their extension."""
    ext = Path(filename).suffix.lower().lstrip(".")
    if ext == "pdf":
        return extract_text_from_pdf(file_path), "pdf"
    if ext in {"doc", "docx"}:
        return extract_text_from_docx(file_path), "docx"
    if ext in {"mp3", "wav", "m4a"}:
        return transcribe_audio_gcp(file_path), "audio"
    if ext in {"txt", "md", "csv"}:
        return Path(file_path).read_text(encoding="utf-8", errors="ignore"), ext
    raise RuntimeError(f"Unsupported file type: .{ext or 'unknown'}")

# =====================================================================
# 3. CHUNKING & INDEXING LAYER (LlamaIndex & ChromaDB)
# =====================================================================
def index_data(workspace_id: str, raw_text: str, source_id: str, source_type: str):
    """
    Tokenizes raw text strings into structured metadata nodes, generates
    embeddings automatically, and saves them to the chosen workspace's vector collection.
    """
    require_ai_dependencies()
    if not raw_text.strip():
        return

   
    chroma_collection = chroma_client.get_or_create_collection(name=workspace_id)
    vector_store = ChromaVectorStore(chroma_collection=chroma_collection)
    storage_context = StorageContext.from_defaults(vector_store=vector_store)

   
    document = Document(
        text=raw_text,
        metadata={"source_id": source_id, "source_type": source_type}
    )

    # Token-based semantic chunk splitter (600 tokens chunk size, 60 tokens overlap)
    text_splitter = TokenTextSplitter(chunk_size=600, chunk_overlap=60)
    nodes = text_splitter.get_nodes_from_documents([document])

    
    VectorStoreIndex.from_documents(
        documents=[document],
        storage_context=storage_context,
        transformations=[text_splitter]
    )

# =====================================================================
# 4. CHATBOT INTEGRATION & SYNTHESIS LAYER
# =====================================================================
def get_dynamic_chat_engine(workspace_id: str, app_identity: str, custom_rules: str) -> CondensePlusContextChatEngine:
    """
    Constructs an interactive chatbot instance on the fly. Isolates context tracking 
    to the current workspace, attaches conversational memory, and sets system behavioral profiles 
    dynamically without any hardcoding.
    """
    require_ai_dependencies()

    chroma_collection = chroma_client.get_or_create_collection(name=workspace_id)
    vector_store = ChromaVectorStore(chroma_collection=chroma_collection)
    storage_context = StorageContext.from_defaults(vector_store=vector_store)
    
    
    index = VectorStoreIndex.from_vector_store(vector_store, storage_context=storage_context)
    
  
    

    dynamic_system_prompt = (
        f"You are the advanced, core AI conversational assistant powering the '{app_identity}' platform.\n"
        f"Your absolute mandate is to solve requests by synthesizing clear, professional responses "
        f"utilizing exclusively the background vector blocks retrieved from the user's Second Memory store.\n\n"
        f"Operational Guardrails:\n"
        f"1. Ground every single assertion inside the provided context blocks.\n"
        f"2. Specific instructions for this active session: {custom_rules}\n"
        f"3. If the background knowledge-base context blocks do not possess the necessary raw information "
        f"to satisfy the query, cleanly state that the secure workspace indices do not currently contain "
        f"records corresponding to that topic."
    )
    
   
    chat_engine = CondensePlusContextChatEngine.from_defaults(
        retriever=index.as_retriever(similarity_top_k=4),
        system_prompt=dynamic_system_prompt,
        verbose=False
    )
    
    return chat_engine


@app.get("/")
def serve_app():
    return FileResponse(PROJECT_ROOT / "second-brain.html")


@app.get("/landing")
def serve_landing():
    return FileResponse(PROJECT_ROOT / "index.html")


@app.get("/api/health")
def health():
    collections = chroma_client.list_collections() if chroma_client else []
    return {
        "ok": True,
        "workspace": DEFAULT_WORKSPACE,
        "ai_dependencies_ready": AI_IMPORT_ERROR is None,
        "ai_dependency_error": str(AI_IMPORT_ERROR) if AI_IMPORT_ERROR else None,
        "llm_configured": bool(GROQ_API_KEY),
        "gcp_configured": os.path.exists(GCP_KEY_PATH),
        "collection_count": len(collections),
    }


@app.post("/api/chat")
def chat(req: ChatRequest):
    message = req.message.strip()
    if not message:
        raise HTTPException(status_code=400, detail="Message cannot be empty.")
    if AI_IMPORT_ERROR is not None:
        return {
            "answer": f"The frontend is connected, but the backend AI dependencies are missing: {AI_IMPORT_ERROR}. Run pip install -r requirements.txt.",
            "workspace_id": req.workspace_id,
        }
    if not GROQ_API_KEY:
        return {
            "answer": "The frontend is connected, but GROQ_API_KEY is not configured on the backend yet.",
            "workspace_id": req.workspace_id,
        }

    try:
        engine = get_dynamic_chat_engine(req.workspace_id, req.app_identity, req.custom_rules)
        response = engine.chat(message)
        return {"answer": str(response), "workspace_id": req.workspace_id}
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc))


@app.post("/api/ingest/text")
def ingest_text(req: TextIngestRequest):
    if not req.text.strip():
        raise HTTPException(status_code=400, detail="Text cannot be empty.")
    try:
        index_data(req.workspace_id, req.text, req.source_id, req.source_type)
        return {"indexed": True, "source_id": req.source_id, "workspace_id": req.workspace_id}
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc))


@app.post("/api/ingest/url")
def ingest_url(req: UrlIngestRequest):
    try:
        raw_text = scrape_url_to_markdown(req.url)
        index_data(req.workspace_id, raw_text, req.url, "url")
        return {"indexed": True, "source_id": req.url, "workspace_id": req.workspace_id}
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc))


@app.post("/api/ingest/files")
async def ingest_files(workspace_id: str = DEFAULT_WORKSPACE, files: List[UploadFile] = File(...)):
    if not files:
        raise HTTPException(status_code=400, detail="Upload at least one file.")

    results = []
    with tempfile.TemporaryDirectory() as temp_dir:
        for upload in files:
            temp_path = Path(temp_dir) / Path(upload.filename).name
            temp_path.write_bytes(await upload.read())
            try:
                raw_text, source_type = extract_text_from_upload(str(temp_path), upload.filename)
                index_data(workspace_id, raw_text, upload.filename, source_type)
                results.append({
                    "filename": upload.filename,
                    "indexed": True,
                    "source_type": source_type,
                })
            except Exception as exc:
                results.append({
                    "filename": upload.filename,
                    "indexed": False,
                    "error": str(exc),
                })

    return {"workspace_id": workspace_id, "files": results}

if __name__ == "__main__":
    print("--- Second Memory AI Engine API: http://127.0.0.1:8000 ---")
    uvicorn.run(app, host=os.getenv("HOST", "127.0.0.1"), port=int(os.getenv("PORT", "8000")))
