import os
import requests
from bs4 import BeautifulSoup
import pypdf
import docx
from datetime import datetime
import re
from dateutil import parser as date_parser

# Optional imports - graceful fallback if not installed
try:
    import qdrant_client
    HAS_QDRANT = True
except ImportError:
    HAS_QDRANT = False
    print("⚠️ Warning: qdrant_client not installed. Vector database features unavailable.")

try:
    from llama_index.core import VectorStoreIndex, StorageContext, Settings
    from llama_index.vector_stores.qdrant import QdrantVectorStore
    from llama_index.llms.groq import Groq
    from llama_index.core.chat_engine import CondensePlusContextChatEngine
    from llama_index.embeddings.huggingface import HuggingFaceEmbedding
    from llama_index.core.node_parser import TokenTextSplitter
    from llama_index.core.schema import Document
    HAS_LLAMA_INDEX = True
except ImportError:
    HAS_LLAMA_INDEX = False
    print("⚠️ Warning: llama_index not installed. AI chat features unavailable.")

try:
    from google.cloud import speech
    HAS_GOOGLE_CLOUD = True
except ImportError:
    HAS_GOOGLE_CLOUD = False
    print("⚠️ Warning: google-cloud-speech not installed. Speech-to-text unavailable.")

# =====================================================================
# 1. INITIALIZATION & STORAGE ENGINE SETUP
# =====================================================================
_ai_layer_dir = os.path.dirname(os.path.abspath(__file__))

# Only initialize llama_index if available
if HAS_LLAMA_INDEX:
    Settings.llm = Groq(
        model="llama-3.3-70b-versatile",
        api_key=os.getenv("GROQ_API_KEY"),
        temperature=0.2
    )

    # Use cloud Inference API on Vercel to stay within the 250MB package limit
    if os.getenv("VERCEL") or os.getenv("HF_TOKEN"):
        from llama_index.embeddings.huggingface import HuggingFaceInferenceAPIEmbedding
        print("🌐 Vercel environment detected. Using Hugging Face Inference API for Embeddings (Stateless)")
        Settings.embed_model = HuggingFaceInferenceAPIEmbedding(
            model_name="sentence-transformers/all-MiniLM-L6-v2",
            token=os.getenv("HF_TOKEN")
        )
    else:
        print("💾 Local environment detected. Using Local Hugging Face Embeddings")
        Settings.embed_model = HuggingFaceEmbedding(model_name="all-MiniLM-L6-v2")

# Initialize Qdrant Client (gracefully falls back to local disk if cloud credentials missing)
qdrant_client_instance = None
if HAS_QDRANT:
    QDRANT_URL = os.getenv("QDRANT_URL")
    QDRANT_API_KEY = os.getenv("QDRANT_API_KEY")
    if QDRANT_URL and QDRANT_API_KEY:
        print(f"🌐 Connecting to Qdrant Cloud at: {QDRANT_URL}")
        qdrant_client_instance = qdrant_client.QdrantClient(
            url=QDRANT_URL,
            api_key=QDRANT_API_KEY
        )
    else:
        local_db_path = os.getenv("DATABASE_PATH") or os.path.join(_ai_layer_dir, "qdrant_db")
        print(f"💾 QDRANT_URL/KEY not found. Fallback to local disk storage at: {local_db_path}")
        qdrant_client_instance = qdrant_client.QdrantClient(path=local_db_path)


# Resolve GCP key path relative to THIS file's directory so it works
# regardless of the working directory api.py is launched from
_script_dir = os.path.dirname(os.path.abspath(__file__))
_gcp_env = os.getenv("GOOGLE_APPLICATION_CREDENTIALS", "")
if _gcp_env and os.path.exists(_gcp_env):
    GCP_KEY_PATH = _gcp_env
elif os.path.exists(os.path.join(_script_dir, "gcp_key.json")):
    GCP_KEY_PATH = os.path.join(_script_dir, "gcp_key.json")
else:
    GCP_KEY_PATH = _gcp_env or os.path.join(_script_dir, "gcp_key.json")

# Always set the env var to the resolved path so GCP client picks it up
os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = GCP_KEY_PATH
if not os.path.exists(GCP_KEY_PATH):
    print(f"⚠️ Warning: GCP credentials file missing at '{GCP_KEY_PATH}'.")
else:
    print(f"✓ GCP credentials loaded from '{GCP_KEY_PATH}'.")

# =====================================================================
# 2. FILE INGESTION PARSERS (Multi-Modal Data Layer)
# =====================================================================
def extract_text_from_pdf(file_path: str) -> str:
    """Extracts raw text content from a local PDF file."""
    try:
        text = ""
        with open(file_path, "rb") as f:
            reader = pypdf.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text.strip()
    except Exception as e:
        raise RuntimeError(f"Failed to parse PDF file: {str(e)}")

def extract_text_from_docx(file_path: str) -> str:
    """Extracts raw text content from a local Word document (.docx)."""
    try:
        doc = docx.Document(file_path)
        full_text = [paragraph.text for paragraph in doc.paragraphs]
        return "\n".join(full_text).strip()
    except Exception as e:
        raise RuntimeError(f"Failed to parse DOCX file: {str(e)}")

def scrape_url_to_markdown(url: str) -> str:
    """Fetches a webpage, cleans non-content tags, and structures pure text."""
    try:
        headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"}
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.text, "html.parser")
        for element in soup(["script", "style", "nav", "footer", "header", "aside"]):
            element.decompose()
            
        lines = (line.strip() for line in soup.get_text().splitlines())
        chunks = (phrase for line in lines for phrase in line.split("  "))
        clean_text = "\n".join(chunk for chunk in chunks if chunk)
        
        return f"# Source: {url}\n\n{clean_text}"
    except Exception as e:
        raise RuntimeError(f"BeautifulSoup scraping failed: {str(e)}")

def transcribe_audio_gcp(audio_file_path: str) -> str:
    """Transcribes an audio file via Google Cloud Speech-to-Text API."""
    if not HAS_GOOGLE_CLOUD:
        raise RuntimeError("google-cloud-speech not installed. Run: pip install google-cloud-speech")
    if not os.path.exists(GCP_KEY_PATH):
        raise RuntimeError(f"GCP credentials not found at '{GCP_KEY_PATH}'. Check GOOGLE_APPLICATION_CREDENTIALS in your .env")
    try:
        client = speech.SpeechClient()
        with open(audio_file_path, "rb") as f:
            audio_content = f.read()
        audio  = speech.RecognitionAudio(content=audio_content)
        config = speech.RecognitionConfig(
            encoding=speech.RecognitionConfig.AudioEncoding.WEBM_OPUS,
            language_code="en-US",
            enable_automatic_punctuation=True,
        )
        # Try WEBM_OPUS first (browser recordings), fall back to LINEAR16
        try:
            response = client.recognize(config=config, audio=audio)
        except Exception:
            config = speech.RecognitionConfig(
                encoding=speech.RecognitionConfig.AudioEncoding.LINEAR16,
                sample_rate_hertz=16000,
                language_code="en-US",
                enable_automatic_punctuation=True,
            )
            response = client.recognize(config=config, audio=audio)
        transcript = " ".join([r.alternatives[0].transcript for r in response.results])
        return transcript if transcript else "[Empty Transcription]"
    except Exception as e:
        raise RuntimeError(f"GCP Speech-to-Text transcription failed: {str(e)}")

def extract_dates_from_text(text: str) -> list:
    """
    Extracts all date references from text with support for:
    - Absolute dates: "2024-12-15", "12/15/2024", "December 15, 2024"
    - Relative dates: "last Tuesday", "next month", "this quarter"
    - Quarter notation: "Q1 2026", "Q2 2025"
    - Month/year only: "March 2026", "Jan '25"
    - Relative references: "3 days ago", "2 weeks from now"
    
    Returns: Sorted list of unique datetime objects
    """
    from datetime import datetime, timedelta
    import calendar
    
    dates = []
    current_date = datetime.now()
    
    # Define weekday names and month names for regex
    weekdays = r'(?:Monday|Tuesday|Wednesday|Thursday|Friday|Saturday|Sunday|Mon|Tue|Wed|Thu|Fri|Sat|Sun)'
    months = r'(?:January|February|March|April|May|June|July|August|September|October|November|December|Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)'
    
    # Pattern 1: Absolute dates (YYYY-MM-DD, MM/DD/YYYY, etc)
    absolute_patterns = [
        r'\d{4}-\d{1,2}-\d{1,2}',  # YYYY-MM-DD
        r'\d{1,2}/\d{1,2}/\d{4}',  # MM/DD/YYYY
        r'(?:' + months + r') \d{1,2},? \d{4}',  # "January 5, 2026"
        r'\d{1,2} (?:' + months + r') \d{4}',  # "5 January 2026"
    ]
    
    for pattern in absolute_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            try:
                parsed_date = date_parser.parse(match.group())
                dates.append(parsed_date)
            except:
                pass
    
    # Pattern 2: Quarter notation (Q1 2026, Q2 2025, etc)
    quarter_pattern = r'Q([1-4])\s*(?:of\s+)?(\d{4}|\d{2})'
    for match in re.finditer(quarter_pattern, text, re.IGNORECASE):
        quarter = int(match.group(1))
        year_str = match.group(2)
        year = int(year_str) if len(year_str) == 4 else 2000 + int(year_str)
        
        # Start of quarter (Q1 = Jan 1, Q2 = Apr 1, etc)
        month = (quarter - 1) * 3 + 1
        dates.append(datetime(year, month, 1))
    
    # Pattern 3: Month/Year notation ("March 2026", "Jan '25")
    month_year_pattern = r'(?:' + months + r')\s*[\']*(\d{4}|\d{2})'
    for match in re.finditer(month_year_pattern, text, re.IGNORECASE):
        try:
            year_str = match.group(1)
            year = int(year_str) if len(year_str) == 4 else 2000 + int(year_str)
            parsed = date_parser.parse(match.group())
            dates.append(datetime(year, parsed.month, 1))
        except:
            pass
    
    # Pattern 4: Relative dates with offsets ("3 days ago", "2 weeks from now")
    relative_pattern = r'(\d+)\s+(day|week|month|year)s?\s+(ago|from\s+now|in\s+the\s+future)'
    for match in re.finditer(relative_pattern, text, re.IGNORECASE):
        amount = int(match.group(1))
        unit = match.group(2).lower()
        direction = match.group(3).lower()
        
        if 'ago' in direction:
            amount = -amount
        
        try:
            if unit == 'day':
                delta_date = current_date + timedelta(days=amount)
            elif unit == 'week':
                delta_date = current_date + timedelta(weeks=amount)
            elif unit == 'month':
                # Add months by calculating target year/month
                target_month = current_date.month + amount
                target_year = current_date.year + (target_month - 1) // 12
                target_month = ((target_month - 1) % 12) + 1
                delta_date = current_date.replace(year=target_year, month=target_month)
            elif unit == 'year':
                delta_date = current_date.replace(year=current_date.year + amount)
            
            dates.append(delta_date)
        except:
            pass
    
    # Pattern 5: Last/Next specific days ("last Tuesday", "next Monday")
    # weekdays uses a non-capturing (?:...) group — wrap in capturing parens for group(2)
    last_next_pattern = r'(last|next|this)\s+(Monday|Tuesday|Wednesday|Thursday|Friday|Saturday|Sunday|Mon|Tue|Wed|Thu|Fri|Sat|Sun)'
    for match in re.finditer(last_next_pattern, text, re.IGNORECASE):
        direction = match.group(1).lower()
        day_name = match.group(2).lower()
        
        # Map day name to weekday number (0=Monday, 6=Sunday)
        day_map = {
            'monday': 0, 'mon': 0,
            'tuesday': 1, 'tue': 1,
            'wednesday': 2, 'wed': 2,
            'thursday': 3, 'thu': 3,
            'friday': 4, 'fri': 4,
            'saturday': 5, 'sat': 5,
            'sunday': 6, 'sun': 6,
        }
        
        target_weekday = day_map.get(day_name.lower())
        if target_weekday is not None:
            try:
                current_weekday = current_date.weekday()
                days_ahead = target_weekday - current_weekday
                
                if direction == 'last':
                    if days_ahead >= 0:
                        days_ahead -= 7
                    delta_date = current_date + timedelta(days=days_ahead)
                elif direction == 'next':
                    if days_ahead <= 0:
                        days_ahead += 7
                    delta_date = current_date + timedelta(days=days_ahead)
                else:  # 'this'
                    if days_ahead < 0:
                        days_ahead += 7
                    delta_date = current_date + timedelta(days=days_ahead)
                
                dates.append(delta_date)
            except:
                pass
    
    # Pattern 6: Season notation (Spring 2026, Summer 2025)
    season_pattern = r'(Spring|Summer|Fall|Autumn|Winter)\s+(\d{4}|\d{2})'
    season_months = {
        'spring': 3,  # March
        'summer': 6,  # June
        'fall': 9,    # September
        'autumn': 9,  # September
        'winter': 12, # December
    }
    for match in re.finditer(season_pattern, text, re.IGNORECASE):
        season = match.group(1).lower()
        year_str = match.group(2)
        year = int(year_str) if len(year_str) == 4 else 2000 + int(year_str)
        
        if season in season_months:
            dates.append(datetime(year, season_months[season], 1))
    
    # Remove duplicates, sort, and return
    unique_dates = list(set(dates))
    return sorted(unique_dates)


# =====================================================================
# 2.5 PRIORITY SCORING LAYER (Options 2 + 3 combined)
# =====================================================================
# Priority levels: 1=HIGH (red), 2=MEDIUM (yellow), 3=LOW (green)
# Stored in ChromaDB metadata as "priority" (int) and "priority_score" (float)
# Dynamic re-scoring happens via /api/nodes which reads query_count from metadata

PRIORITY_URGENCY_WORDS = [
    "urgent", "asap", "deadline", "due date", "overdue", "critical", "must",
    "immediately", "emergency", "action required", "by tomorrow", "by end of",
    "expires", "expiring", "time-sensitive", "priority", "do not miss"
]
PRIORITY_MEDIUM_WORDS = [
    "meeting", "task", "follow up", "follow-up", "todo", "to-do", "to do",
    "schedule", "review", "check", "remind", "next steps", "action item",
    "plan", "agenda", "note", "notes"
]

def score_priority(text: str, source_type: str, ingestion_timestamp: str) -> dict:
    """
    Scores document priority using Option 2 (content-based) + Option 3 (recency boost).
    Returns dict with priority (1/2/3) and signals for transparency.
    
    Priority 1 = HIGH (red)   — urgent content or very recent
    Priority 2 = MEDIUM (yellow) — task/meeting/action content
    Priority 3 = LOW (green)  — reference material
    """
    text_lower = text[:3000].lower()  # Only scan first 3000 chars for speed

    # Option 2: Content-based signals
    urgency_hits = sum(1 for w in PRIORITY_URGENCY_WORDS if w in text_lower)
    medium_hits  = sum(1 for w in PRIORITY_MEDIUM_WORDS  if w in text_lower)

    # Source type boosts
    source_boost = 0
    if source_type in ("LOG", "AUDIO"):
        source_boost = 1   # Conversations/recordings get a bump
    elif source_type in ("PDF", "DOCX"):
        source_boost = 0   # Formal docs are neutral
    elif source_type == "WEB":
        source_boost = -1  # Web articles default to low

    # Option 3: Recency boost — docs ingested in last 24h get +1 urgency signal
    try:
        ingested_at = datetime.fromisoformat(ingestion_timestamp)
        hours_old = (datetime.now() - ingested_at).total_seconds() / 3600
        recency_boost = 2 if hours_old < 24 else (1 if hours_old < 72 else 0)
    except Exception:
        recency_boost = 0

    # Composite score
    raw_score = urgency_hits * 3 + medium_hits + source_boost + recency_boost

    if raw_score >= 5 or urgency_hits >= 2:
        priority = 1  # HIGH
    elif raw_score >= 2 or medium_hits >= 2:
        priority = 2  # MEDIUM
    else:
        priority = 3  # LOW

    return {
        "priority": priority,
        "priority_score": float(raw_score),
        "priority_urgency_hits": urgency_hits,
        "priority_medium_hits": medium_hits,
        "priority_recency_boost": recency_boost,
    }


def recompute_priority(existing_meta: dict, query_count: int) -> int:
    """
    Option 3: Dynamic re-prioritisation based on usage.
    Frequently queried nodes climb in priority over time.
    Called by /api/nodes when building the node list.
    """
    base = existing_meta.get("priority", 3)
    score = existing_meta.get("priority_score", 0.0)

    # Every 5 queries boosts the score by 1 point
    usage_boost = query_count // 5
    adjusted_score = score + usage_boost

    if adjusted_score >= 5:
        return 1
    elif adjusted_score >= 2:
        return 2
    else:
        return base


# =====================================================================
# 2.7 SEMANTIC MEMORY CLUSTERING LAYER
# =====================================================================
def get_embed_model():
    if HAS_LLAMA_INDEX and hasattr(Settings, "embed_model") and Settings.embed_model:
        return Settings.embed_model
    from llama_index.embeddings.huggingface import HuggingFaceEmbedding
    return HuggingFaceEmbedding(model_name="all-MiniLM-L6-v2")

def generate_cluster_name(text: str) -> str:
    """Uses Groq LLM to generate a human-like 2-3 word topic name for a cluster."""
    import groq as groq_sdk
    try:
        client = groq_sdk.Groq(api_key=os.getenv("GROQ_API_KEY"))
        prompt = (
            "Analyze the following document text and suggest a very short, high-level semantic category or topic name "
            "representing this document (e.g., 'Project Orion', 'Financial Statements', 'AI & Machine Learning', "
            "'Personal Health', 'Meeting Minutes'). "
            "The name must be exactly 2 to 4 words. Respond with only the name, no punctuation or intro.\n\n"
            f"Document text excerpt:\n{text[:1000]}"
        )
        completion = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=20,
        )
        res = completion.choices[0].message.content or ""
        clean = res.strip().replace('"', '').replace("'", "")
        return clean if clean else "General Knowledge"
    except Exception as e:
        print(f"Error generating cluster name: {e}")
        return "General Knowledge"

def assign_cluster(workspace_id: str, raw_text: str, source_id: str) -> dict:
    """
    Finds or creates a semantic cluster for the new document.
    Returns a dict with "cluster_id" and "cluster_name".
    """
    if qdrant_client_instance is None:
        return {"cluster_id": "cluster_default", "cluster_name": "General"}
        
    try:
        if not qdrant_client_instance.collection_exists(collection_name=workspace_id):
            cluster_name = generate_cluster_name(raw_text)
            return {"cluster_id": "cluster_1", "cluster_name": cluster_name}
            
        scroll_results = qdrant_client_instance.scroll(
            collection_name=workspace_id,
            with_payload=True,
            with_vectors=True,
            limit=10000
        )
        points = scroll_results[0]
        
        if not points:
            cluster_name = generate_cluster_name(raw_text)
            return {"cluster_id": "cluster_1", "cluster_name": cluster_name}
            
        import numpy as np
        embed_model = get_embed_model()
        new_embedding = np.array(embed_model.get_text_embedding(raw_text[:1000]))
        new_embedding_norm = new_embedding / np.linalg.norm(new_embedding)
        
        best_sim = -1.0
        best_meta = None
        
        # Calculate similarity with existing document chunks
        for p in points:
            emb = p.vector
            payload = p.payload or {}
            meta = payload.get("metadata", payload)
            if meta.get("source_id") == "chat_log.jsonl":
                continue
            if emb is None:
                continue
            emb_arr = np.array(emb)
            norm_val = np.linalg.norm(emb_arr)
            if norm_val == 0:
                continue
            emb_norm = emb_arr / norm_val
            sim = float(np.dot(new_embedding_norm, emb_norm))
            if sim > best_sim:
                best_sim = sim
                best_meta = meta
                
        # If similarity is strong (>= 0.62), group in the same cluster
        if best_sim >= 0.62 and best_meta and best_meta.get("cluster_id"):
            return {
                "cluster_id": best_meta.get("cluster_id"),
                "cluster_name": best_meta.get("cluster_name", "General Knowledge")
            }
            
        # Otherwise, create a new cluster
        existing_cluster_ids = set()
        for p in points:
            payload = p.payload or {}
            meta = payload.get("metadata", payload)
            cid = meta.get("cluster_id")
            if cid:
                existing_cluster_ids.add(cid)
                
        cluster_nums = []
        for cid in existing_cluster_ids:
            try:
                cluster_nums.append(int(cid.split("_")[1]))
            except:
                pass
        next_num = max(cluster_nums) + 1 if cluster_nums else 1
        new_cluster_id = f"cluster_{next_num}"
        new_cluster_name = generate_cluster_name(raw_text)
        
        return {
            "cluster_id": new_cluster_id,
            "cluster_name": new_cluster_name
        }
    except Exception as e:
        print(f"Error in assign_cluster: {e}")
        return {"cluster_id": "cluster_default", "cluster_name": "General Knowledge"}

def ensure_database_clustering(workspace_id: str):
    """
    Ensures all documents in the Qdrant collection are clustered.
    If any document is missing a cluster_id, it performs a batch clustering
    on all documents and updates their metadata in Qdrant.
    """
    if qdrant_client_instance is None:
        return
        
    try:
        if not qdrant_client_instance.collection_exists(collection_name=workspace_id):
            return
            
        scroll_results = qdrant_client_instance.scroll(
            collection_name=workspace_id,
            with_payload=True,
            with_vectors=True,
            limit=10000
        )
        points = scroll_results[0]
        
        if not points:
            return
            
        needs_clustering = False
        for p in points:
            payload = p.payload or {}
            meta = payload.get("metadata", payload)
            if meta.get("source_id") == "chat_log.jsonl":
                continue
            if "cluster_id" not in meta or not meta["cluster_id"]:
                needs_clustering = True
                break
                
        if not needs_clustering:
            return
            
        print(f"🔄 Missing clusters detected in Qdrant workspace '{workspace_id}'. Running batch clustering...")
        
        docs_by_source = {}
        for p in points:
            payload = p.payload or {}
            meta = payload.get("metadata", payload)
            doc_text = payload.get("text", "")
            emb = p.vector
            sid = meta.get("source_id")
            if not sid or sid == "chat_log.jsonl":
                continue
            if sid not in docs_by_source:
                docs_by_source[sid] = {
                    "ids": [],
                    "texts": [],
                    "embeddings": []
                }
            docs_by_source[sid]["ids"].append(p.id)
            docs_by_source[sid]["texts"].append(doc_text)
            if emb is not None:
                docs_by_source[sid]["embeddings"].append(emb)
            
        if not docs_by_source:
            return
            
        import numpy as np
        doc_ids = list(docs_by_source.keys())
        doc_embeddings = []
        for sid in doc_ids:
            embs = [e for e in docs_by_source[sid]["embeddings"] if e is not None]
            if embs:
                doc_embeddings.append(np.mean(embs, axis=0))
            else:
                embed_model = get_embed_model()
                txt = docs_by_source[sid]["texts"][0][:1000] if docs_by_source[sid]["texts"] else sid
                doc_embeddings.append(np.array(embed_model.get_text_embedding(txt)))
                
        clusters = {}
        cluster_names = {}
        cluster_count = 0
        assigned = {}
        
        for i, sid_i in enumerate(doc_ids):
            if sid_i in assigned:
                continue
            cluster_count += 1
            cluster_num = cluster_count
            clusters[cluster_num] = [sid_i]
            assigned[sid_i] = cluster_num
            
            first_text = docs_by_source[sid_i]["texts"][0] if docs_by_source[sid_i]["texts"] else sid_i
            cluster_names[cluster_num] = generate_cluster_name(first_text)
            
            emb_i = doc_embeddings[i]
            norm_i = np.linalg.norm(emb_i)
            emb_i_norm = emb_i / norm_i if norm_i > 0 else emb_i
            
            for j, sid_j in enumerate(doc_ids):
                if sid_j in assigned:
                    continue
                emb_j = doc_embeddings[j]
                norm_j = np.linalg.norm(emb_j)
                emb_j_norm = emb_j / norm_j if norm_j > 0 else emb_j
                sim = float(np.dot(emb_i_norm, emb_j_norm))
                if sim >= 0.62:
                    clusters[cluster_num].append(sid_j)
                    assigned[sid_j] = cluster_num
                    
        # Update point payloads in Qdrant
        updated_count = 0
        for p in points:
            payload = p.payload or {}
            meta = dict(payload.get("metadata", payload))
            sid = meta.get("source_id")
            if not sid or sid == "chat_log.jsonl":
                continue
            cnum = assigned.get(sid)
            if cnum:
                meta["cluster_id"] = f"cluster_{cnum}"
                meta["cluster_name"] = cluster_names[cnum]
                
                # Update payload in Qdrant
                qdrant_client_instance.set_payload(
                    collection_name=workspace_id,
                    payload={"metadata": meta},
                    points=[p.id]
                )
                updated_count += 1
                
        print(f"✅ Successfully clustered {len(doc_ids)} documents into {cluster_count} semantic clusters in Qdrant (updated {updated_count} points).")
            
    except Exception as e:
        print(f"Error during batch Qdrant database clustering: {e}")

# =====================================================================
# 3. CHUNKING & INDEXING LAYER (LlamaIndex & Qdrant)
# =====================================================================
def index_data(workspace_id: str, raw_text: str, source_id: str, source_type: str):
    """
    Tokenizes raw text into chunks, stamps EVERY chunk with shared metadata
    (source_id, source_type, first_mentioned_date), then stores in Qdrant.
    LlamaIndex does NOT propagate custom metadata to child nodes by default —
    we do it explicitly here so timeline/calendar queries work correctly.
    """
    if not raw_text.strip():
        return

    # Use QdrantVectorStore to manage vectors and collection automatically
    vector_store = QdrantVectorStore(client=qdrant_client_instance, collection_name=workspace_id)
    storage_context = StorageContext.from_defaults(vector_store=vector_store)

    # Extract temporal metadata from full document text
    extracted_dates = extract_dates_from_text(raw_text)
    first_date = extracted_dates[0].isoformat() if extracted_dates else None

    # Compute priority before building metadata
    ingestion_ts = datetime.now().isoformat()
    priority_info = score_priority(raw_text, source_type, ingestion_ts)

    # Find/assign semantic cluster
    cluster_info = assign_cluster(workspace_id, raw_text, source_id)

    # Shared metadata stamped on every chunk
    shared_metadata = {
        "source_id": source_id,
        "source_type": source_type,
        "ingestion_timestamp": ingestion_ts,
        "first_mentioned_date": first_date,
        "mentioned_dates_count": len(extracted_dates),
        "query_count": 0,
        "cluster_id": cluster_info["cluster_id"],
        "cluster_name": cluster_info["cluster_name"],
        **priority_info,
    }

    document = Document(text=raw_text, metadata=shared_metadata)

    # Split then explicitly copy metadata to every child node
    text_splitter = TokenTextSplitter(chunk_size=600, chunk_overlap=60)
    nodes = text_splitter.get_nodes_from_documents([document])
    for node in nodes:
        node.metadata.update(shared_metadata)

    # Index pre-built nodes (avoids double-splitting)
    VectorStoreIndex(
        nodes=nodes,
        storage_context=storage_context,
    )

# =====================================================================
# 3.5 TIMELINE SYNTHESIS LAYER
# =====================================================================
def generate_timeline(workspace_id: str, start_date: str = None, end_date: str = None) -> str:
    """
    Synthesizes a chronological timeline from indexed documents in Qdrant.
    Returns a markdown-formatted timeline of events mentioned in the workspace.
    
    Args:
        workspace_id: The workspace to query
        start_date: ISO format date string (YYYY-MM-DD) to filter from
        end_date: ISO format date string (YYYY-MM-DD) to filter until
    """
    try:
        if qdrant_client_instance is None or not qdrant_client_instance.collection_exists(collection_name=workspace_id):
            return "No timeline data available. No documents have been indexed yet."
        
        # Retrieve all documents/metadata from Qdrant via scroll
        scroll_results = qdrant_client_instance.scroll(
            collection_name=workspace_id,
            with_payload=True,
            with_vectors=False,
            limit=10000
        )
        points = scroll_results[0]
        
        if not points:
            return "No timeline data available. No documents have been indexed yet."
        
        # Deduplicate by source_id — keep earliest date and first excerpt per file
        seen = {}  # source_id -> event dict
        for p in points:
            metadata = p.payload.get("metadata", p.payload) or {}
            doc_text = p.payload.get("text", "")
            first_date_str = metadata.get("first_mentioned_date")
            if not first_date_str:
                continue
            try:
                event_date = datetime.fromisoformat(first_date_str)

                # Apply date range filter if provided
                if start_date:
                    start = datetime.fromisoformat(start_date)
                    if event_date < start:
                        continue
                if end_date:
                    end = datetime.fromisoformat(end_date)
                    if event_date > end:
                        continue

                sid = metadata.get("source_id", "Unknown")
                if sid not in seen or event_date < seen[sid]["date"]:
                    seen[sid] = {
                        "date": event_date,
                        "source_id": sid,
                        "source_type": metadata.get("source_type", "Unknown"),
                        "ingestion_time": metadata.get("ingestion_timestamp", "Unknown"),
                        "excerpt": doc_text[:150] + "..." if len(doc_text) > 150 else doc_text
                    }
            except:
                pass

        timeline_events = list(seen.values())
        
        if not timeline_events:
            return f"No events found in timeline (range: {start_date} to {end_date})"
        
        # Sort chronologically
        timeline_events.sort(key=lambda x: x["date"])
        
        # Format as markdown
        timeline_md = "# Timeline\n\n"
        for event in timeline_events:
            timeline_md += f"## {event['date'].strftime('%B %d, %Y')}\n"
            timeline_md += f"**Source:** `{event['source_type']}` ({event['source_id']})\n"
            timeline_md += f"**Indexed:** {event['ingestion_time']}\n"
            timeline_md += f"\n> {event['excerpt']}\n\n"
        
        return timeline_md
        
    except Exception as e:
        return f"Error generating timeline: {str(e)}"

# =====================================================================
# 4. CHATBOT INTEGRATION & SYNTHESIS LAYER
# =====================================================================
def query_workspace(workspace_id: str, message: str, app_identity: str, custom_rules: str) -> str:
    """
    Retrieve top-k chunks from Qdrant via embedding similarity,
    then call the Groq SDK *directly* — bypassing LlamaIndex's HTTP
    wrapper which was returning gzip-compressed bytes as raw text.
    """
    import groq as groq_sdk

    if qdrant_client_instance is None or not qdrant_client_instance.collection_exists(collection_name=workspace_id):
        return "I could not find that in your indexed documents."

    # ── 1. Embed the query and retrieve relevant chunks ──────────────
    embed_model = get_embed_model()
    query_embedding = embed_model.get_text_embedding(message)

    results = qdrant_client_instance.search(
        collection_name=workspace_id,
        query_vector=query_embedding,
        limit=8
    )
    
    # Group context chunks by their cluster name
    clustered_context = {}
    for r in results:
        payload = r.payload or {}
        meta = payload.get("metadata", payload) or {}
        doc = payload.get("text", "")
        cname = meta.get("cluster_name", "General Knowledge")
        sid = meta.get("source_id", "Unknown Document")
        if cname not in clustered_context:
            clustered_context[cname] = []
        clustered_context[cname].append(f"- Document: {sid}\n  Content: {doc}")

    context_str = ""
    for cname, items in clustered_context.items():
        context_str += f"=== Topic Cluster: {cname} ===\n"
        context_str += "\n".join(items) + "\n\n"

    # ── 2. Build prompt ───────────────────────────────────────────────
    system_msg = (
        f"You are a warm, highly conversational, and intelligent personal memory assistant for the '{app_identity}' knowledge base. "
        f"{custom_rules} "
        f"Your voice is natural, engaging, friendly, and fully human-like (not structured like a machine). "
        f"Answer the user's question using the provided context. The context chunks are grouped by their semantic topics (clusters). "
        f"Synthesize the details across these topics naturally. Avoid dry openings like 'Based on the context...' and speak directly. "
        f"If the query is a simple greeting or general conversational query (e.g. 'hello', 'who are you', 'how does this work'), "
        f"respond in a warm, welcoming way directly without requiring specific document search results. "
        f"If the user is asking about something that is not in the context, do not respond with a rigid error. "
        f"Instead, explain in a natural, polite human way that you couldn't find information on that specific topic "
        f"in their documents, and suggest what related items or terms they could check."
    )
    user_msg = f"Clustered context from indexed documents:\n{context_str}\n\nQuestion: {message}"

    # ── 3. Call Groq SDK directly — clean text response guaranteed ────
    client = groq_sdk.Groq(api_key=os.getenv("GROQ_API_KEY"))
    completion = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {"role": "system", "content": system_msg},
            {"role": "user",   "content": user_msg},
        ],
        temperature=0.4,
        max_tokens=1024,
    )
    answer = completion.choices[0].message.content or ""
    return answer.strip() or "I could not find that in your indexed documents."


# Kept for import compatibility — delegates to query_workspace
def get_dynamic_chat_engine(workspace_id: str, app_identity: str, custom_rules: str):
    raise NotImplementedError("Replaced by query_workspace(). Update api.py.")

def detect_and_handle_timeline_request(query: str, workspace_id: str) -> tuple:
    """
    Detects if a query is asking for a timeline and returns (is_timeline_request, response).
    If it is, generates and returns the timeline directly.
    """
    EXPLICIT_TIMELINE_PHRASES = [
        "show timeline", "show me the timeline", "generate timeline",
        "give me a timeline", "display timeline", "build a timeline",
        "create a timeline", "chronological list", "chronological view",
        "show all events", "list all events", "show milestones",
        "project timeline", "full timeline", "complete timeline",
    ]
    query_lower = query.lower().strip()
    is_timeline = any(phrase in query_lower for phrase in EXPLICIT_TIMELINE_PHRASES)
    
    if not is_timeline:
        return False, None
    
    # Try to extract date range if specified
    start_date = None
    end_date = None
    
    # Simple extraction - could be more sophisticated
    if "from" in query_lower and "to" in query_lower:
        # Very basic parsing - in production use more robust date extraction
        parts = query_lower.split("from")[1].split("to")
        try:
            start_date = date_parser.parse(parts[0].strip()).date().isoformat()
            end_date = date_parser.parse(parts[1].strip()).date().isoformat()
        except:
            pass
    
    timeline = generate_timeline(workspace_id, start_date, end_date)
    return True, timeline

if __name__ == "__main__":
    print("--- Second Memory AI Engine Core: Operational and Clean ---")